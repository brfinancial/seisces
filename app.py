import math
import re
import textwrap
from datetime import datetime
from io import BytesIO

import matplotlib.pyplot as plt
import pdfplumber
import streamlit as st
from docx import Document

# =============== CONFIGURAÇÃO DAS PERGUNTAS ===============

QUESTIONS = {
"Caráter": [
{"id": "caracter_hist_atrasos_desc", "type": "text",
"text": "Descreva, com sinceridade, o histórico da empresa em relação a atrasos com bancos, fornecedores e impostos nos últimos 24 meses."},
{"id": "caracter_hist_atrasos_nota", "type": "scale",
"text": "De 0 a 10, qual nota você daria para o cumprimento de prazos de pagamento da empresa? (0 = péssimo, 10 = sempre em dia)"},
{"id": "caracter_restricoes_desc", "type": "text",
"text": "Existem protestos, apontamentos em Serasa/Cartórios ou restrições em nome da empresa ou dos sócios? Descreva."},
{"id": "caracter_restricoes_nota", "type": "scale",
"text": "De 0 a 10, quão limpa você considera a situação cadastral da empresa e dos sócios? (0 = muito suja, 10 = totalmente limpa)"},
{"id": "caracter_reputacao_desc", "type": "text",
"text": "Como você descreveria a reputação da empresa junto a fornecedores, clientes e parceiros?"},
{"id": "caracter_reputacao_nota", "type": "scale",
"text": "De 0 a 10, qual nota você daria para a reputação da empresa no mercado?"},
],
"Capacidade": [
{"id": "capacidade_fluxo_desc", "type": "text",
"text": "Descreva como está hoje o fluxo de caixa da empresa (entradas, saídas, aperto em determinados períodos etc.)."},
{"id": "capacidade_fluxo_nota", "type": "scale",
"text": "De 0 a 10, qual nota você daria para a capacidade atual da empresa de gerar caixa para pagar dívidas?"},
{"id": "capacidade_faturamento_desc", "type": "text",
"text": "Explique como está o faturamento dos últimos 12 meses (crescendo, caindo, estável)."},
{"id": "capacidade_faturamento_nota", "type": "scale",
"text": "De 0 a 10, quão confortável você está com o nível atual de faturamento para suportar novas dívidas?"},
{"id": "capacidade_endividamento_desc", "type": "text",
"text": "Descreva o nível de endividamento atual (bancos, factorings, fornecedores, impostos)."},
{"id": "capacidade_endividamento_nota", "type": "scale",
"text": "De 0 a 10, considerando tudo, qual nota você daria para a capacidade da empresa de assumir mais crédito sem se complicar?"},
],
"Capital": [
{"id": "capital_estrutura_desc", "type": "text",
"text": "Descreva a estrutura financeira da empresa: possui reservas, capital próprio, patrimônio, bens em nome da empresa?"},
{"id": "capital_reservas_nota", "type": "scale",
"text": "De 0 a 10, qual nota você daria para o nível de reserva financeira e capital próprio da empresa?"},
{"id": "capital_patrimonio_desc", "type": "text",
"text": "Quais são os principais bens e ativos relevantes em nome da empresa (imóveis, máquinas, veículos etc.)?"},
{"id": "capital_patrimonio_nota", "type": "scale",
"text": "De 0 a 10, quão robusto você considera o patrimônio da empresa em relação ao tamanho do negócio?"},
{"id": "capital_resiliencia_desc", "type": "text",
"text": "Como a empresa costuma reagir a crises (perda de clientes, queda de faturamento, aumento de custo)?"},
{"id": "capital_resiliencia_nota", "type": "scale",
"text": "De 0 a 10, qual a capacidade da empresa de suportar períodos difíceis sem deixar de pagar suas obrigações?"},
],
"Colateral": [
{"id": "colateral_bens_desc", "type": "text",
"text": "Que garantias a empresa poderia oferecer em uma operação de crédito (imóveis, veículos, máquinas, recebíveis)?"},
{"id": "colateral_bens_nota", "type": "scale",
"text": "De 0 a 10, quão fortes e líquidos você considera esses bens como garantia?"},
{"id": "colateral_recebiveis_desc", "type": "text",
"text": "A empresa possui carteira de recebíveis (duplicatas, boletos, cartões, contratos) que poderia ser usada como garantia? Descreva."},
{"id": "colateral_recebiveis_nota", "type": "scale",
"text": "De 0 a 10, qual a qualidade desses recebíveis (prazo, risco de inadimplência, concentração em poucos clientes)?"},
{"id": "colateral_avales_desc", "type": "text",
"text": "Os sócios estariam dispostos a dar garantias pessoais (aval, fiança) se necessário? Descreva."},
{"id": "colateral_avales_nota", "type": "scale",
"text": "De 0 a 10, quão confortável você considera a estrutura de garantias que a empresa conseguiria montar hoje?"},
],
"Condições": [
{"id": "condicoes_setor_desc", "type": "text",
"text": "Descreva como está o momento do setor em que a empresa atua (expansão, crise, concorrência forte etc.)."},
{"id": "condicoes_setor_nota", "type": "scale",
"text": "De 0 a 10, quão favoráveis são as condições do setor para a empresa hoje?"},
{"id": "condicoes_economia_desc", "type": "text",
"text": "Como a situação econômica geral (juros, inflação, demanda) tem impactado a empresa?"},
{"id": "condicoes_economia_nota", "type": "scale",
"text": "De 0 a 10, quão confortável é o cenário econômico atual para assumir crédito?"},
{"id": "condicoes_operacao_desc", "type": "text",
"text": "Qual seria a finalidade principal do crédito (capital de giro, investimento, alongamento de dívida etc.)?"},
{"id": "condicoes_operacao_nota", "type": "scale",
"text": "De 0 a 10, quão coerente você considera a tomada de crédito com a realidade atual da empresa?"},
],
"Conglomerado": [
{"id": "conglomerado_grupo_desc", "type": "text",
"text": "A empresa faz parte de um grupo econômico? Descreva rapidamente as empresas relacionadas e relações entre elas."},
{"id": "conglomerado_grupo_nota", "type": "scale",
"text": "De 0 a 10, quanto você considera que o grupo econômico fortalece a empresa (em vez de enfraquecer)?"},
{"id": "conglomerado_socios_desc", "type": "text",
"text": "Descreva o perfil dos sócios e da gestão (experiência, envolvimento no dia a dia, alinhamento)."},
{"id": "conglomerado_socios_nota", "type": "scale",
"text": "De 0 a 10, qual nota você daria para a qualidade da gestão e dos sócios da empresa?"},
{"id": "conglomerado_controles_desc", "type": "text",
"text": "A empresa possui controles internos, contabilidade organizada, relatórios financeiros e acompanhamento de indicadores? Descreva."},
{"id": "conglomerado_controles_nota", "type": "scale",
"text": "De 0 a 10, quão estruturada você considera a governança e os controles da empresa?"},
]
}

# =============== HEURÍSTICAS AUXILIARES ===============

POSITIVE_WORDS = [
"em dia", "pontual", "sem atrasos", "sem atraso", "crescente", "crescendo",
"estável", "aumentando", "melhorando", "reservas", "lucro", "lucrativo",
"sem restrição", "sem protesto", "limpo", "organizado", "estruturado",
"controle", "governança", "bom relacionamento", "boa reputação"
]

NEGATIVE_WORDS = [
"atraso", "atrasos", "inadimplência", "inadimplente", "protesto", "protestos",
"serasa", "restrição", "restrições", "crise", "queda", "caindo", "dificuldade",
"aperto", "negativo", "prejuízo", "endividado", "endividamento alto",
"sem reserva", "sem garantia", "desorganizado", "bagunça"
]


def risk_color(percent: float) -> str:
if percent >= 80:
return "🟢 Verde (risco baixo)"
elif 60 <= percent < 80:
return "🟡 Amarelo (risco moderado)"
elif 40 <= percent < 60:
return "🟠 Laranja (risco elevado)"
else:
return "🔴 Vermelho (risco crítico)"


def analyze_text_block(text: str, category: str) -> str:
t = text.lower()
pos = sum(t.count(w) for w in POSITIVE_WORDS)
neg = sum(t.count(w) for w in NEGATIVE_WORDS)

if pos == 0 and neg == 0 and not t.strip():
return "Não houve informações qualitativas suficientes declaradas nessa dimensão para um diagnóstico mais fino."

if pos > neg:
base = "As respostas qualitativas indicam tendência mais positiva nessa dimensão, com alguns pontos que jogam a favor da empresa."
elif neg > pos:
base = "As respostas qualitativas sugerem presença de fragilidades relevantes nessa dimensão, exigindo atenção redobrada."
else:
base = "As respostas qualitativas mostram um cenário misto, com fatores positivos e negativos se equilibrando."

if category == "Caráter":
complemento = " Em Caráter, isso se traduz em histórico e postura que impactam diretamente a confiança na empresa."
elif category == "Capacidade":
complemento = " Em Capacidade, essa leitura afeta diretamente a percepção sobre geração de caixa e capacidade de honrar compromissos."
elif category == "Capital":
complemento = " Em Capital, isso reflete o quão preparada a empresa está estruturalmente para suportar choques e imprevistos."
elif category == "Colateral":
complemento = " Em Colateral, o foco é a consistência e qualidade das garantias que poderiam mitigar o risco assumido."
elif category == "Condições":
complemento = " Em Condições, a leitura recai sobre o ambiente externo e a aderência da tomada de crédito ao momento do negócio."
else:  # Conglomerado
complemento = " Em Conglomerado, essa percepção está ligada à força do grupo econômico, gestão e governança."

return base + complemento


# =============== EXTRAÇÃO GENÉRICA DE PDF ===============

def extract_pdf_text(file) -> str:
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
    return text


def parse_br_number(num_str: str):
    try:
        clean = num_str.replace(".", "").replace(",", ".").strip()
        return float(clean)
    except Exception:
        return None


# =============== ANÁLISE SERASA ===============

def analyze_serasa_text(text: str) -> str:
    tl = text.lower()

    protest_value = None
    protest_match = re.search(r"protest[oa]s?.{0,80}?r\$\s*([\d\.\,]+)", text, flags=re.IGNORECASE | re.DOTALL)
    if protest_match:
        protest_value = parse_br_number(protest_match.group(1))

    frases_bom_fornecedor = [
        "não foram encontradas pendências comerciais",
        "não constam pendências comerciais",
        "sem pendências comerciais",
        "sem pendências com fornecedores"
    ]
    good_suppliers = any(frase in tl for frase in frases_bom_fornecedor)
    has_supplier_pendencias = "pendências comerciais" in tl or "pendencias comerciais" in tl

    bank_terms = [" banco ", "bancária", "bancario", "instituição financeira", "instituicoes financeiras",
                  "financeira", "crédito bancário", "operações de crédito", "operacoes de credito"]
    bank_hits = sum(tl.count(t) for t in bank_terms)

    bank_negative = any(p in tl for p in [
        "atraso com bancos", "pendência com instituições financeiras",
        "pendências com instituições financeiras", "crédito bancário em atraso",
        "em atraso com instituições financeiras"
    ])

    tax_terms = ["dívida ativa", "divida ativa", "receita federal", "débito tributário", "debito tributario",
                 "tributário", "tributario", "inss", "fgts", "icms", "iss", "imposto", "tributos"]
    tax_hits = sum(tl.count(t) for t in tax_terms)

    if bank_hits == 0:
        bancos_txt = "O relatório não traz elementos claros sobre endividamento com instituições financeiras; é recomendável validar com DFs e outras fontes."
    elif bank_negative:
        bancos_txt = "Há menções a pendências ou atrasos junto a instituições financeiras, indicando endividamento bancário com sinais de estresse."
    else:
        bancos_txt = "Existem referências a bancos/financeiras, mas sem evidência forte de atraso; o endividamento bancário parece presente, porém sem sinais claros de deterioração."

    if good_suppliers:
        fornecedores_txt = "O relatório indica bom histórico de pagamento a fornecedores na praça, sem pendências comerciais relevantes."
    elif has_supplier_pendencias:
        fornecedores_txt = "Constam pendências comerciais com fornecedores, o que sugere fragilidade na cadeia de pagamentos com a praça."
    else:
        fornecedores_txt = "Não há indicação clara de pendências comerciais com fornecedores; a situação parece neutra ou não detalhada."

    if tax_hits > 0:
        impostos_txt = "Há sinais de envolvimento com temas tributários (dívida ativa, Receita Federal ou débitos de impostos), sugerindo passivos fiscais que devem ser considerados na análise."
    else:
        impostos_txt = "O relatório não evidencia de forma explícita débitos tributários relevantes, ou essas informações não estão claras no texto extraído."

    if protest_value is not None:
        if protest_value >= 50000 and good_suppliers:
            dica_txt = (
                "Observa-se um valor elevado em protestos, mas com bom histórico de pagamento a fornecedores. "
                "Essa combinação, na prática de análise de crédito, costuma indicar concentração de atrasos em "
                "obrigações fiscais ou específicas (como tributos), o que é menos grave do que ruptura direta "
                "da cadeia de fornecedores, embora ainda exija atenção na modelagem da operação."
            )
        elif protest_value >= 50000 and not good_suppliers:
            dica_txt = (
                "O valor de protestos é relevante e não há evidência de bom histórico com fornecedores, "
                "o que aponta para um risco mais sensível de crédito, incluindo possíveis problemas na praça."
            )
        else:
            dica_txt = (
                "Há registro de protestos, mas em valor que não se mostra excessivamente elevado pelo texto capturado. "
                "Ainda assim, é prudente cruzar as informações com balanços, DRE e fluxo de caixa projetado."
            )
    else:
        dica_txt = (
            "Não foi possível identificar com clareza o valor total de protestos no texto extraído. "
            "Sugere-se conferir manualmente o quadro específico de protestos do relatório."
        )

    resumo = (
        "Endividamento com bancos: " + bancos_txt + " "
        "Histórico de pagamento a fornecedores: " + fornecedores_txt + " "
        "Situação de impostos e tributos: " + impostos_txt + " "
        + dica_txt
    )

    return resumo


def serasa_section():
    st.subheader("Análise de Relatório Serasa (PDF) – opcional")

    if "serasa_resumo" not in st.session_state:
        st.session_state["serasa_resumo"] = None

    uploaded = st.file_uploader("Envie o relatório Serasa (PDF):", type=["pdf"], key="serasa_pdf")

    if uploaded is not None:
        if st.button("Analisar relatório Serasa"):
            try:
                text = extract_pdf_text(uploaded)
                resumo = analyze_serasa_text(text)
                st.session_state["serasa_resumo"] = resumo
                st.success("Relatório Serasa analisado com sucesso.")
            except Exception as e:
                st.error(f"Não foi possível ler o PDF do Serasa. Detalhe técnico: {e}")

    if st.session_state["serasa_resumo"]:
        st.markdown("### Resumo da análise do Serasa")
        st.write(st.session_state["serasa_resumo"])

    return st.session_state["serasa_resumo"]


# =============== ANÁLISE SISBACEN / SCR ===============

def analyze_sisbacen_text(text: str) -> str:
    """
    Leitura heurística de um relatório SISBACEN/SCR:
    - Exposição com bancos
    - Presença de atrasos / risco elevado
    - Operações baixadas a prejuízo
    - Humaniza o parecer bancário
    """
    tl = text.lower()

    # Sinais de exposição relevante
    termos_exposicao = [
        "exposição total", "exposicao total", "saldo devedor", "limite contratado",
        "valor total das operações", "operações de crédito", "operacoes de credito",
        "risco total"
    ]
    exp_hits = sum(tl.count(t) for t in termos_exposicao)

    # Sinais de atraso / classificação de risco ruim
    termos_atraso = [
        "em atraso", "vencida", "vencidas", "vencidos", "inadimplência", "inadimplente",
        "atraso superior", "faixa de atraso", "dias de atraso"
    ]
    atraso_hits = sum(tl.count(t) for t in termos_atraso)

    # Classificações de risco típicas (AA, A, B, C, D, E, F, G, H)
    classes_ruins = ["risco e", "risco f", "risco g", "risco h", "classificação e", "classificação f",
                     "classificacao e", "classificacao f", "classificacao g", "classificacao h"]
    risco_ruim_hits = sum(tl.count(c) for c in classes_ruins)

    # Operações baixadas a prejuízo
    termos_prejuizo = [
        "baixa a prejuízo", "baixa a prejuizo", "baixada para prejuízo", "baixada para prejuizo",
        "operações baixadas como prejuízo", "operacoes baixadas como prejuizo"
    ]
    prejuizo_hits = sum(tl.count(t) for t in termos_prejuizo)

    # Tentativa simples de achar algum valor de exposição total
    exp_valor = None
    exp_match = re.search(r"(exposi[cç][aã]o total|risco total|valor total das opera[cç][õo]es).{0,80}?r\$\s*([\d\.\,]+)",
                          text, flags=re.IGNORECASE | re.DOTALL)
    if exp_match:
        exp_valor = parse_br_number(exp_match.group(2))

    # Montagem da análise humanizada

    # 1) Exposição bancária
    if exp_valor is not None:
        if exp_valor < 100000:
            exp_txt = f"A exposição bancária total identificada gira em torno de R$ {exp_valor:,.2f}, em patamar relativamente contido para a maioria das PMEs."
        elif exp_valor < 500000:
            exp_txt = f"A exposição bancária total estimada é de aproximadamente R$ {exp_valor:,.2f}, o que indica uso relevante de linhas bancárias, porém ainda administrável dependendo do faturamento."
        else:
            exp_txt = f"A exposição bancária total aparenta ser elevada, na casa de cerca de R$ {exp_valor:,.2f}, sugerindo alavancagem relevante junto ao sistema financeiro."
    else:
        if exp_hits > 0:
            exp_txt = "O relatório menciona saldos de operações de crédito e exposição com bancos, mas o valor total não pôde ser determinado de forma clara pelo texto extraído."
        else:
            exp_txt = "Não foi possível identificar com clareza o montante de exposição total com bancos; recomenda-se consultar diretamente os quadros de valores do relatório."

    # 2) Comportamento de atraso / risco
    if risco_ruim_hits > 0 or atraso_hits > 5:
        comportamento_txt = (
            "O histórico bancário apresenta sinais de atraso e/ou classificação de risco em faixas mais pressionadas "
            "(como E, F, G ou H), indicando que parte relevante das operações já transitou para um patamar de maior risco."
        )
        perfil_risco = "pressionado/crítico"
    elif atraso_hits > 0:
        comportamento_txt = (
            "Há registros de atraso em algumas operações, mas sem indicação consistente de concentração nas piores faixas "
            "de risco. Ainda assim, é um ponto de atenção na concessão de novos créditos."
        )
        perfil_risco = "moderado"
    else:
        comportamento_txt = (
            "Não foram identificadas referências fortes a atrasos ou níveis de risco críticos, sugerindo um histórico "
            "de relacionamento bancário relativamente bem comportado."
        )
        perfil_risco = "saudável"

    # 3) Operações baixadas a prejuízo
    if prejuizo_hits > 0:
        prejuizo_txt = (
            "Constam operações baixadas a prejuízo, o que indica que, em algum momento, instituições financeiras "
            "tiveram de reconhecer perda efetiva com o tomador. Esse é um sinal relevante e pesa de forma negativa "
            "na análise de crédito, exigindo estruturação mais conservadora das operações e, se possível, apoio em garantias."
        )
        tem_prejuizo = True
    else:
        prejuizo_txt = (
            "Não foram identificadas, no texto extraído, menções claras a operações baixadas a prejuízo, o que reduz "
            "a percepção de histórico de default bancário extremo."
        )
        tem_prejuizo = False

    # 4) Conclusão humanizada (tom bancário x factoring)
    if perfil_risco == "saudável" and not tem_prejuizo:
        conclusao_txt = (
            "De forma geral, o quadro bancário sugere um tomador que utiliza o sistema financeiro de maneira relativamente "
            "organizada, sem sinais contundentes de estresse prolongado. Para fins de fomento/factoring, isso abre espaço "
            "para operações com limites compatíveis ao faturamento, mantendo disciplina de monitoramento."
        )
    elif perfil_risco == "moderado" and not tem_prejuizo:
        conclusao_txt = (
            "O conjunto das informações indica um tomador com relacionamento bancário já um pouco tensionado, "
            "mas ainda recuperável. A recomendação é trabalhar com limites mais enxutos, prazos mais curtos e "
            "cláusulas que permitam rápida reação em caso de piora, usando o fomento mais como ferramenta de "
            "organização do fluxo de caixa do que de alavancagem adicional agressiva."
        )
    else:
        conclusao_txt = (
            "O histórico bancário aponta para um perfil de maior risco, seja pelo acúmulo de atrasos em faixas ruins "
            "de classificação, seja pela presença de operações baixadas a prejuízo. Isso sugere que a empresa já "
            "enfrentou momentos de estresse considerável com bancos. Nessa situação, qualquer concessão de crédito "
            "deve ser pensada de forma muito criteriosa, com foco em operações estruturadas, valores menores, "
            "prazos curtos e, sempre que possível, reforço de garantias ou vinculação direta a recebíveis de boa qualidade."
        )

    resumo = (
        "Exposição com bancos: " + exp_txt + " "
        "Comportamento de atraso e risco: " + comportamento_txt + " "
        "Operações baixadas a prejuízo: " + prejuizo_txt + " "
        + conclusao_txt
    )

    return resumo


def sisbacen_section():
    st.subheader("Análise de Relatório SISBACEN / SCR (PDF) – opcional")

    if "sisbacen_resumo" not in st.session_state:
        st.session_state["sisbacen_resumo"] = None

    uploaded = st.file_uploader("Envie o relatório SISBACEN/SCR (PDF):", type=["pdf"], key="sisbacen_pdf")

    if uploaded is not None:
        if st.button("Analisar relatório SISBACEN/SCR"):
            try:
                text = extract_pdf_text(uploaded)
                resumo = analyze_sisbacen_text(text)
                st.session_state["sisbacen_resumo"] = resumo
                st.success("Relatório SISBACEN/SCR analisado com sucesso.")
            except Exception as e:
                st.error(f"Não foi possível ler o PDF do SISBACEN/SCR. Detalhe técnico: {e}")

    if st.session_state["sisbacen_resumo"]:
        st.markdown("### Resumo da análise de crédito bancário (SISBACEN/SCR)")
        st.write(st.session_state["sisbacen_resumo"])

    return st.session_state["sisbacen_resumo"]


# =============== PARECER EM PROSA ===============

def generate_report(company_name, all_answers, category_scores, overall_percent,
                    sazonalidade_resumo=None, serasa_resumo=None):
def generate_report(
    company_name,
    all_answers,
    category_scores,
    overall_percent,
    sazonalidade_resumo=None,
    serasa_resumo=None,
    sisbacen_resumo=None,
):
wrapper = textwrap.TextWrapper(width=100)
lines = []

lines.append(f"Parecer de Crédito - Empresa: {company_name}")
lines.append("=" * 100)
lines.append(f"Data da análise: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
lines.append(f"Score geral de crédito: {overall_percent:.1f}%")
lines.append(f"Nível de risco: {risk_color(overall_percent)}")
lines.append("")

if sazonalidade_resumo:
lines.append("Resumo de sazonalidade de crédito:")
lines.append(wrapper.fill(sazonalidade_resumo))
lines.append("")

if serasa_resumo:
lines.append("Resumo da análise do relatório Serasa:")
lines.append(wrapper.fill(serasa_resumo))
lines.append("")

    if sisbacen_resumo:
        lines.append("Resumo da análise de crédito bancário (SISBACEN / SCR):")
        lines.append(wrapper.fill(sisbacen_resumo))
        lines.append("")

if overall_percent >= 80:
visao_geral = (
"Na minha avaliação, a empresa apresenta um perfil de crédito globalmente saudável. "
"Os fundamentos de capacidade de pagamento, organização e estrutura de suporte ao crédito "
"aparecem bem posicionados, permitindo uma exposição maior com risco relativamente controlado."
)
elif overall_percent >= 60:
visao_geral = (
"Na minha leitura, a empresa demonstra um perfil de crédito razoável. Há pontos consistentes, "
"mas também algumas vulnerabilidades que sugerem prudência na definição de limites, prazos e "
"eventuais garantias. O crédito é possível, mas deve ser estruturado com critério."
)
elif overall_percent >= 40:
visao_geral = (
"Com base nas respostas fornecidas, o perfil de crédito da empresa apresenta fragilidades "
"significativas. A concessão de crédito deve ser feita com bastante cautela, em valores menores, "
"prazos mais curtos e forte amparo em garantias, até que os pontos críticos sejam endereçados."
)
else:
visao_geral = (
"Pela combinação das informações qualitativas e quantitativas, o perfil atual é de alto risco. "
"Há elementos que indicam baixa capacidade de suportar novas dívidas sem agravamento da situação "
"financeira. Minha opinião é que, neste momento, a empresa deveria priorizar reorganização interna "
"e ajuste de estrutura antes de novas concessões."
)

lines.append(wrapper.fill(visao_geral))
lines.append("")
lines.append("Resumo por dimensão (6 C’s do crédito):")

for cat, data in category_scores.items():
lines.append(f"- {cat}: {data['percent']:.1f}% ({risk_color(data['percent'])})")

lines.append("")
lines.append("Análise qualitativa e opinião por C:")

for cat, questions in QUESTIONS.items():
lines.append("")
lines.append(f"--- {cat.upper()} ---")

cat_percent = category_scores[cat]["percent"]
if cat_percent >= 80:
base_comment = f"Numérica e comparativamente, {cat} aparece como um ponto forte da empresa."
elif cat_percent >= 60:
base_comment = f"Em {cat}, os indicadores mostram um nível aceitável, porém com sinais que merecem acompanhamento."
elif cat_percent >= 40:
base_comment = f"Os resultados em {cat} revelam fragilidades relevantes, que podem se refletir em risco adicional na concessão de crédito."
else:
base_comment = f"Em {cat}, a pontuação indica um ponto crítico, que tende a pressionar negativamente a decisão de crédito."

lines.append(wrapper.fill(base_comment))

cat_text_block = ""
for q in questions:
if q["type"] == "text":
ans = all_answers.get(q["id"], "")
if ans:
cat_text_block += " " + ans

comentario_qualitativo = analyze_text_block(cat_text_block, cat)
lines.append("")
lines.append(wrapper.fill(comentario_qualitativo))

if cat == "Caráter":
recomendacao = (
"Recomendo reforçar o histórico de pontualidade, regularizar eventuais restrições e manter "
"uma postura transparente com credores e fornecedores, pois isso sustenta a confiança no longo prazo."
)
elif cat == "Capacidade":
recomendacao = (
"É importante aprimorar planejamento de fluxo de caixa, monitorar de perto endividamento e proteger "
"a margem operacional, garantindo que novas dívidas sejam suportáveis."
)
elif cat == "Capital":
recomendacao = (
"Fortalecer capital próprio, patrimônio e reservas aumenta a resiliência da empresa e reduz a "
"sensibilidade a choques de mercado ou perda de clientes."
)
elif cat == "Colateral":
recomendacao = (
"Estruturar garantias mais consistentes – seja em bens de boa liquidez, seja em recebíveis de qualidade – "
"melhora substancialmente a atratividade da empresa para operações de crédito."
)
elif cat == "Condições":
recomendacao = (
"Vale alinhar o uso do crédito ao momento setorial e macroeconômico, priorizando operações que apoiem "
"ajuste de estrutura ou crescimento sustentável, e não apenas o fechamento de buracos de curto prazo."
)
else:  # Conglomerado
recomendacao = (
"Aperfeiçoar governança, clareza nas relações entre empresas do grupo, qualidade da gestão e registros "
"contábeis ajuda a reduzir opacidade e transmitir mais segurança a quem concede crédito."
)

lines.append("")
lines.append("Recomendação nesta dimensão:")
lines.append(wrapper.fill(recomendacao))

return "\n".join(lines)


# =============== WORD EM MEMÓRIA (PARA DOWNLOAD) ===============

def generate_word_doc_bytes(company_name, report_text, overall_percent):
safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', company_name) or "Empresa"
filename = f"Parecer_Credito_{safe_name}.docx"

doc = Document()
doc.add_heading("Parecer de Crédito Empresarial", level=1)
doc.add_paragraph(f"Empresa: {company_name}")
doc.add_paragraph(f"Score geral de crédito: {overall_percent:.1f}%")
doc.add_paragraph(f"Nível de risco: {risk_color(overall_percent)}")
doc.add_paragraph(f"Data da análise: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
doc.add_paragraph("")

doc.add_heading("Conclusão e opinião da IA", level=2)

for bloco in report_text.split("\n\n"):
bloco = bloco.strip()
if bloco:
doc.add_paragraph(bloco)

bio = BytesIO()
doc.save(bio)
bio.seek(0)
return filename, bio


# =============== SAZONALIDADE (STREAMLIT) ===============

def sazonalidade_section():
st.subheader("Sazonalidade de crédito")

possui_sazonalidade = st.radio(
"O negócio possui sazonalidade relevante ao longo do ano?",
["Não", "Sim"],
horizontal=True
)

if possui_sazonalidade == "Não":
        st.info("Não há sazonalidade específica nesse setor segundo as informações fornecidas.")
        return "Não há sazonalidade específica nesse setor segundo as informações fornecidas.", None
        msg = "Não há sazonalidade específica nesse setor segundo as informações fornecidas."
        st.info(msg)
        return msg, None

setor = st.text_input(
"Informe o setor de atuação da empresa (ex.: fantasias, varejo, agro, serviços etc.):",
""
).lower()

pico_vendas = st.selectbox(
"Mês de pico de vendas/demanda:",
options=list(range(1, 13)),
format_func=lambda m: ["Jan", "Fev", "Mar", "Abr", "Mai", "Jun", "Jul", "Ago", "Set", "Out", "Nov", "Dez"][m - 1]
)

estrategia_default = 2
if any(p in setor for p in ["agro", "agrícola", "agronegócio", "soja", "safra", "grãos"]):
estrategia_default = 2  # durante

st.write("Em relação a esse pico de demanda, quando costuma ser mais adequado conceder crédito?")
estrategia = st.radio(
"Janela preferencial de crédito:",
[
"Alguns meses antes (preparação / formação de estoque)",
"Durante o próprio pico (ex.: agro no verão)",
"Logo depois do pico (pós-safra / pós-temporada)"
],
index=estrategia_default - 1
)

if estrategia.startswith("Alguns meses antes"):
shift = -2
elif estrategia.startswith("Durante"):
shift = 0
else:
shift = 1

centro_credito = pico_vendas + shift
if centro_credito < 1:
centro_credito += 12
elif centro_credito > 12:
centro_credito -= 12

meses = list(range(1, 13))
nomes_meses = ["Jan", "Fev", "Mar", "Abr", "Mai", "Jun",
"Jul", "Ago", "Set", "Out", "Nov", "Dez"]

sigma = 2.0
valores_brutos = []
for m in meses:
dist_direta = abs(m - centro_credito)
dist_circular = min(dist_direta, 12 - dist_direta)
valor = math.exp(-(dist_circular ** 2) / (2 * sigma ** 2))
valores_brutos.append(valor)

max_valor = max(valores_brutos) if valores_brutos else 1
valores_percentuais = [(v / max_valor) * 100 for v in valores_brutos]

fig, ax = plt.subplots(figsize=(8, 4))
ax.plot(meses, valores_percentuais, marker="o")
ax.set_xticks(meses)
ax.set_xticklabels(nomes_meses)
ax.set_ylim(0, 110)
ax.set_xlabel("Meses do ano")
ax.set_ylabel("Atratividade de aprovação de crédito (%)")
ax.set_title("Sazonalidade recomendada de crédito")
ax.grid(True)

st.pyplot(fig)

if shift == -2:
janela = "alguns meses ANTES do pico de vendas"
elif shift == 0:
janela = "no próprio pico de vendas"
else:
janela = "logo DEPOIS do pico de vendas"

nome_mes_pico = nomes_meses[pico_vendas - 1]
nome_mes_centro = nomes_meses[centro_credito - 1]
resumo = (
f"Para o setor informado ({setor if setor else 'não especificado'}), a análise considera que o "
f"pico de demanda ocorre em {nome_mes_pico}. A janela ótima de crédito foi ajustada para {janela}, "
f"com maior atratividade concentrada em {nome_mes_centro} e meses próximos."
)

st.caption(resumo)
return resumo, fig


# =============== SERASA: EXTRAÇÃO E ANÁLISE ===============

def extract_pdf_text(file) -> str:
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
    return text


def parse_br_number(num_str: str):
    try:
        clean = num_str.replace(".", "").replace(",", ".").strip()
        return float(clean)
    except Exception:
        return None


def analyze_serasa_text(text: str) -> str:
    """
    Análise heurística simples de relatório Serasa:
    - Endividamento com bancos
    - Pagamento a fornecedores
    - Situação de impostos (via inferência)
    - Dica: protestos altos + bons fornecedores -> provável peso de impostos
    """
    tl = text.lower()

    # --- Protestos (tentativa de capturar valor total) ---
    protest_value = None
    protest_match = re.search(r"protest[oa]s?.{0,80}?r\$\s*([\d\.\,]+)", text, flags=re.IGNORECASE | re.DOTALL)
    if protest_match:
        protest_value = parse_br_number(protest_match.group(1))

    # Bom histórico de fornecedores?
    frases_bom_fornecedor = [
        "não foram encontradas pendências comerciais",
        "não constam pendências comerciais",
        "sem pendências comerciais",
        "sem pendências com fornecedores"
    ]
    good_suppliers = any(frase in tl for frase in frases_bom_fornecedor)

    # Indícios gerais de pendências comerciais
    has_supplier_pendencias = "pendências comerciais" in tl or "pendencias comerciais" in tl

    # Bancos / financeiras
    bank_terms = [" banco ", "bancária", "bancario", "instituição financeira", "instituicoes financeiras",
                  "financeira", "crédito bancário", "operacões de crédito", "operações de crédito"]
    bank_hits = sum(tl.count(t) for t in bank_terms)

    bank_negative = any(p in tl for p in ["atraso com bancos", "pendência com instituições financeiras",
                                          "pendências com instituições financeiras",
                                          "crédito bancário em atraso", "em atraso com instituições financeiras"])

    # Impostos / tributos
    tax_terms = ["dívida ativa", "divida ativa", "receita federal", "débito tributário", "debito tributario",
                 "tributário", "tributario", "inss", "fgts", "icms", "iss", "imposto", "tributos"]
    tax_hits = sum(tl.count(t) for t in tax_terms)

    # --- Montagem da análise ---

    # Bancos
    if bank_hits == 0:
        bancos_txt = "O relatório não traz elementos claros sobre endividamento com instituições financeiras; é recomendável validar com DFs e outras fontes."
    elif bank_negative:
        bancos_txt = "Há menções a pendências ou atrasos junto a instituições financeiras, indicando endividamento bancário com sinais de estresse."
    else:
        bancos_txt = "Existem referências a bancos/financeiras, mas sem evidência forte de atraso; o endividamento bancário parece presente, porém sem sinais claros de deterioração."

    # Fornecedores
    if good_suppliers:
        fornecedores_txt = "O relatório indica bom histórico de pagamento a fornecedores na praça, sem pendências comerciais relevantes."
    elif has_supplier_pendencias:
        fornecedores_txt = "Constam pendências comerciais com fornecedores, o que sugere fragilidade na cadeia de pagamentos com a praça."
    else:
        fornecedores_txt = "Não há indicação clara de pendências comerciais com fornecedores; a situação parece neutra ou não detalhada."

    # Impostos / tributos
    if tax_hits > 0:
        impostos_txt = "Há sinais de envolvimento com temas tributários (dívida ativa, Receita Federal ou débitos de impostos), sugerindo passivos fiscais que devem ser considerados na análise."
    else:
        impostos_txt = "O relatório não evidencia de forma explícita débitos tributários relevantes, ou essas informações não estão claras no texto extraído."

    # Heurística: protestos altos + bons fornecedores -> provável imposto
    dica_txt = ""
    if protest_value is not None:
        if protest_value >= 50000 and good_suppliers:
            dica_txt = (
                "Observa-se um valor elevado em protestos, mas com bom histórico de pagamento a fornecedores. "
                "Essa combinação, na prática de análise de crédito, costuma indicar concentração de atrasos em "
                "obrigações fiscais ou específicas (como tributos), o que é menos grave do que ruptura direta "
                "da cadeia de fornecedores, embora ainda exija atenção na modelagem da operação."
            )
        elif protest_value >= 50000 and not good_suppliers:
            dica_txt = (
                "O valor de protestos é relevante e não há evidência de bom histórico com fornecedores, "
                "o que aponta para um risco mais sensível de crédito, incluindo possíveis problemas na praça."
            )
        else:
            dica_txt = (
                "Há registro de protestos, mas em valor que não se mostra excessivamente elevado pelo texto capturado. "
                "Ainda assim, é prudente cruzar as informações com balanços, DRE e fluxo de caixa projetado."
            )
    else:
        dica_txt = (
            "Não foi possível identificar com clareza o valor total de protestos no texto extraído. "
            "Sugere-se conferir manualmente o quadro específico de protestos do relatório."
        )

    resumo = (
        "Endividamento com bancos: " + bancos_txt + " "
        "Histórico de pagamento a fornecedores: " + fornecedores_txt + " "
        "Situação de impostos e tributos: " + impostos_txt + " "
        + dica_txt
    )

    return resumo


def serasa_section():
    st.subheader("Análise de Relatório Serasa (PDF) – opcional")

    if "serasa_resumo" not in st.session_state:
        st.session_state["serasa_resumo"] = None

    uploaded = st.file_uploader("Envie o relatório Serasa (PDF):", type=["pdf"], key="serasa_pdf")

    if uploaded is not None:
        if st.button("Analisar relatório Serasa"):
            try:
                text = extract_pdf_text(uploaded)
                resumo = analyze_serasa_text(text)
                st.session_state["serasa_resumo"] = resumo
                st.success("Relatório Serasa analisado com sucesso.")
            except Exception as e:
                st.error(f"Não foi possível ler o PDF do Serasa. Detalhe técnico: {e}")

    if st.session_state["serasa_resumo"]:
        st.markdown("### Resumo da análise do Serasa")
        st.write(st.session_state["serasa_resumo"])

    return st.session_state["serasa_resumo"]


# =============== APP STREAMLIT ===============

def main():
st.set_page_config(page_title="IA de Crédito Empresarial - BRF", layout="wide")

st.title("IA de Diagnóstico de Crédito Empresarial")
    st.write("Baseada nos 6 C’s do crédito – versão BR Financial com parecer opinativo, sazonalidade e leitura de Serasa (PDF).")
    st.write("Baseada nos 6 C’s do crédito – versão BR Financial com parecer opinativo, sazonalidade, Serasa e SISBACEN/SCR.")

company_name = st.text_input("Nome da empresa analisada:", "")

all_answers = {}
category_scores = {}

st.header("Questionário – 6 C’s do crédito")

for category, questions in QUESTIONS.items():
with st.expander(category, expanded=False):
cat_score = 0.0
cat_max = 0.0

for q in questions:
if q["type"] == "text":
ans = st.text_area(q["text"], key=q["id"])
else:
ans = st.slider(q["text"], 0.0, 10.0, 5.0, 0.5, key=q["id"])
cat_score += ans
cat_max += 10.0
all_answers[q["id"]] = ans

cat_percent = (cat_score / cat_max) * 100 if cat_max > 0 else 0
category_scores[category] = {
"score": cat_score,
"max": cat_max,
"percent": cat_percent
}
st.markdown(f"**Score parcial de {category}: {cat_percent:.1f}% ({risk_color(cat_percent)})**")

st.markdown("---")
sazonalidade_resumo, _ = sazonalidade_section()

st.markdown("---")
serasa_resumo = serasa_section()

    st.markdown("---")
    sisbacen_resumo = sisbacen_section()

st.markdown("---")
if st.button("Gerar parecer e documento Word"):
company_name_use = company_name if company_name else "Empresa Não Informada"

total_score = sum(c["score"] for c in category_scores.values())
total_max = sum(c["max"] for c in category_scores.values())
overall_percent = (total_score / total_max) * 100 if total_max > 0 else 0

report = generate_report(
company_name_use,
all_answers,
category_scores,
overall_percent,
sazonalidade_resumo=sazonalidade_resumo,
            serasa_resumo=serasa_resumo
            serasa_resumo=serasa_resumo,
            sisbacen_resumo=sisbacen_resumo,
)

st.subheader("Parecer de crédito")
st.text(report)

filename, word_bytes = generate_word_doc_bytes(company_name_use, report, overall_percent)

st.download_button(
label="Baixar parecer em Word",
data=word_bytes,
file_name=filename,
mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


if __name__ == "__main__":
main()
